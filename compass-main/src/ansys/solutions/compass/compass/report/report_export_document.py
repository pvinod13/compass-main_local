# ©2022, ANSYS Inc. Unauthorized use, distribution or duplication is prohibited.

# ==================================================== [Imports] ==================================================== #

import docx
from docx.shared import Cm
from docx2pdf import convert
import pythoncom
import os
import tempfile

from ansys.report.sdk.report import Report
from ansys.report.sdk.report_image import ReportImage
from ansys.report.sdk.report_table import ReportTable

from ansys.solutions.compass.compass.report.report_dto import Image, Paragraph
from ansys.solutions.compass.compass.report.report_section import ReportSection
from ansys.solutions.compass.compass.report.report_table import ReportTable

# =================================================== [Functions] =================================================== #


def export_document(report_dto):
    exportPdf = False # To set to true when ready to use ansys.report.sdk
    file = tempfile.NamedTemporaryFile().name

    if exportPdf:
        pdf_file = file + ".pdf"
        with Report() as report:
            part = report.add_part('Report')
            __render_pdf_section(part, report_dto.introduction_section)
            __render_pdf_section(part, report_dto.materials_section)
            __render_pdf_section(part, report_dto.maxwell_analysis_section)
            report.export_pdf(pdf_file)
    else:
        docx_file = file + ".docx"
        document = docx.Document()
        document.add_heading("Report", level=0)
        __render_docx_section(document, report_dto.introduction_section)
        __render_docx_section(document, report_dto.materials_section)
        __render_docx_section(document, report_dto.maxwell_analysis_section)
        document.save(docx_file)
        pdf_file = file + ".pdf"
        convert(docx_file, pdf_file, pythoncom.CoInitialize())
    return pdf_file


def __render_docx_section(document, section):
    heading = document.add_heading(section.title, section.level)
    heading.style = document.styles["Heading " + str(section.level)]
    if section.contents:
        for content in section.contents:
            if isinstance(content, Paragraph):
                document.add_paragraph(content.Text)
            elif isinstance(content, ReportSection):
                __render_docx_section(document, content)
            elif isinstance(content, ReportTable):
                __render_docx_table(document, content)
            elif isinstance(content, Image):
                __render_docx_image(document, content)
            else:
                raise NotImplementedError()


def __render_docx_table(document, report_table):
    data = report_table.content
    column_count = len(data)
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Medium Grid 2 Accent 1"

    header = table.rows[0].cells
    for index, key in enumerate(data):
        header[index].text = str(key)

    row_count = len(list(data.values())[0])
    for _ in range(row_count):
        table.add_row()

    for row_index in range(1, row_count + 1):
        row_cells = table.rows[row_index].cells
        for column_index in range(column_count):
            row_cells[column_index].text = str(list(data.values())[column_index][row_index - 1])


def __render_docx_image(document, image):
    document.add_picture(os.path.join(os.path.abspath(__file__), image.Path))
    picture = document.add_picture(os.path.join(os.path.abspath(__file__), image.Path))
    picture.width = Cm(image.Width)
    picture.height = Cm(image.Height)


def __render_pdf_section(pdf_part, section):
    pdf_section = pdf_part.add_section(section.title)
    if section.contents:
        for content in section.contents:
            if isinstance(content, Paragraph):
                pdf_section.add_texts(content.Text)
            elif isinstance(content, ReportSection):
                __render_pdf_section(pdf_part, content)
            elif isinstance(content, ReportTable):
                __render_pdf_table(pdf_section, content)
            elif isinstance(content, Image):
                __render_pdf_image(pdf_section, content)
            else:
                raise NotImplementedError()


def __render_pdf_table(pdf_section, inputs):
    rows = [list(inputs.keys())]
    column_count = len(inputs.content)
    row_count = len(list(inputs.content.values())[0])

    for row_index in range(row_count):
        row = []
        for column_index in range(column_count):
            row.append(str(list(inputs.content.values())[column_index][row_index]))
        rows.append(row)

    column_width_list = column_count * [4]

    pdf_table = ReportTable(
        rows=rows,
        column_width_list=column_width_list,
        caption=inputs.Caption
    )
    pdf_section.add_table(pdf_table)


def __render_pdf_image(pdf_section, image):
    report_image = ReportImage(
        image_file_path=image.Path,
        width = image.Width,
        caption = image.Caption
    )
    pdf_section.add_image(report_image)
